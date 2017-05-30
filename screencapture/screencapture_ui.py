from tkinter import *
from PIL import ImageGrab
from docx import Document
from docx.shared import Inches
import sys


document = Document()
counter = 0
filename_prefix = 'demo'
directory = 'C:'
docx_name = "sample.docx"
default_filename = 'sample.docx'


def clear_text():
    e.delete(0, 'end')

def capture_screen():
    global counter
    global docx_name
    global document
    comments = e.get()
    snapshot = ImageGrab.grab()
    str_counter = str(counter)
    save_path = directory + '\\python\\' + filename_prefix + str_counter + '.jpg'
    snapshot.save(save_path)
    p = document.add_paragraph()
    r = p.add_run()
    r.add_picture(save_path, width=Inches(7.0))
    r.add_text(comments)
    counter = counter + 1
    document.save(directory + '\\python\\' + docx_name + '.docx')

def printtext():
    global e
    global counter
    global docx_name
    global document
    
    filename_input = T.get("1.0", 'end-1c')
    
    if filename_input != docx_name:
        docx_name = filename_input
        document = Document()
        print("updated file name")
        #print(docx_name)
        
    elif filename_input == docx_name:
        print ("writing in same file")
       # print (docx_name)
    root.iconify()
    capture_screen()
    

root = Tk()
root.title('CAPTURE')

## comment area tag
comment = Label(root, text="Comments:")
comment.pack(side=LEFT)

## comment area of utility
e = Entry(root)
e.pack(side=LEFT)
e.focus_set()

## filename tag for filename text area
tag = Label(root, text="Filename:")
tag.pack(side=LEFT)

## created filename text area
T = Text(root, height=1, width=20)
T.pack(side=LEFT)
T.insert(END, default_filename)

## capture button for taking screenshots
capture_button = Button(root,text='capture',command=printtext)
capture_button.pack(side='left')

## clear the comments 
clear_button = Button(root, text='clear Text', command=clear_text)
clear_button.pack(side='left')
root.mainloop()
