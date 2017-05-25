from pynput.keyboard import Key, Listener
from PIL import ImageGrab
from docx import Document
from docx.shared import Inches
import sys
#from Xlib.display import Display

document = Document()
counter = 0
filename_prefix = 'demo'
directory = 'C:'
docx_prefix = sys.argv[1]
docx_counter = 1
new_file_key = Key.shift
screenshot_key = Key.print_screen

def on_press(key):
    if key == screenshot_key:
       print('{0} pressed'.format(
        key))
    if key == new_file_key :
        print('creating new file')

def on_release(key):

    global counter
    global docx_counter
    global document
    if key == screenshot_key :
       print('{0} release'.format(
        key))
       snapshot = ImageGrab.grab()
       str_counter = str(counter)
       str_docx_counter = str(docx_counter)
       save_path = directory + '\\python\\' + filename_prefix + str_counter + '.jpg'
       snapshot.save(save_path)
       p = document.add_paragraph()
       r = p.add_run()
       r.add_picture(save_path, width=Inches(7.0))
       
       counter = counter + 1
       document.save(directory + '\\python\\' + docx_prefix + '_' + str_docx_counter + '.docx')
    if key == new_file_key :
       docx_counter = docx_counter + 1
       document = Document()
       
    if key == Key.esc:
        # Stop listener
        return False

# Collect events until released
with Listener(
        on_press=on_press,
        on_release=on_release) as listener:
    listener.join()
